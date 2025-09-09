"""
Document processing module for the AI Research Assistant.

This module handles extraction of text from various document formats.
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any
import tempfile
import os

# Document processing libraries
import PyPDF2
import pypdf
from docx import Document

from utils.config import validate_file_size, validate_file_type

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles extraction of text from various document formats."""
    
    def __init__(self, max_file_size_mb: int = 50):
        """
        Initialize the document processor.
        
        Args:
            max_file_size_mb (int): Maximum file size in MB
        """
        self.max_file_size_mb = max_file_size_mb
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Extract text from a document file.
        
        Args:
            file_path (str): Path to the document file
            filename (str): Original filename
            
        Returns:
            Dict[str, Any]: Dictionary containing extracted text and metadata
        """
        result = {
            'success': False,
            'text': '',
            'filename': filename,
            'file_type': Path(filename).suffix.lower(),
            'error': None,
            'metadata': {}
        }
        
        try:
            # Validate file type
            if not validate_file_type(filename):
                result['error'] = f"Unsupported file type: {result['file_type']}"
                return result
            
            # Validate file size
            if not validate_file_size(file_path, self.max_file_size_mb):
                result['error'] = f"File too large (max {self.max_file_size_mb}MB)"
                return result
            
            # Extract text based on file type
            if result['file_type'] == '.pdf':
                text, metadata = self._extract_pdf_text(file_path)
            elif result['file_type'] == '.docx':
                text, metadata = self._extract_docx_text(file_path)
            elif result['file_type'] == '.txt':
                text, metadata = self._extract_txt_text(file_path)
            else:
                result['error'] = f"Unsupported file type: {result['file_type']}"
                return result
            
            if text and text.strip():
                result['success'] = True
                result['text'] = text
                result['metadata'] = metadata
                logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            else:
                result['error'] = "No text could be extracted from the document. The file might be scanned, empty, or corrupted."
            
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _extract_pdf_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using PyPDF2 with pypdf fallback.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            tuple[str, Dict[str, Any]]: Extracted text and metadata
        """
        text = ""
        metadata = {'pages': 0, 'method': 'unknown'}
        
        # Validate file first
        if not os.path.exists(file_path):
            raise Exception(f"PDF file not found: {file_path}")
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise Exception("PDF file is empty (0 bytes)")
        
        # Check if it's a valid PDF by reading header
        try:
            with open(file_path, 'rb') as file:
                header = file.read(5)
                if not header.startswith(b'%PDF-'):
                    raise Exception("File is not a valid PDF (missing PDF header)")
        except Exception as e:
            if "PDF header" in str(e):
                raise e
            raise Exception(f"Cannot read file: {e}")
        
        # Try PyPDF2 first
        pypdf2_error = None
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise Exception("PDF has no pages")
                
                metadata['pages'] = len(pdf_reader.pages)
                metadata['method'] = 'PyPDF2'
                
                extracted_pages = 0
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                            extracted_pages += 1
                    except Exception as page_error:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                        continue
                
                # If we got some text, consider it successful
                if text.strip():
                    logger.info(f"PyPDF2: Extracted text from {extracted_pages}/{len(pdf_reader.pages)} pages")
                    return text.strip(), metadata
                
        except Exception as e:
            pypdf2_error = str(e)
            logger.warning(f"PyPDF2 failed: {e}. Trying pypdf...")
        
        # Fallback to pypdf
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise Exception("PDF has no pages")
                
                metadata['pages'] = len(pdf_reader.pages)
                metadata['method'] = 'pypdf'
                
                extracted_pages = 0
                text = ""  # Reset text for pypdf attempt
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                            extracted_pages += 1
                    except Exception as page_error:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                        continue
                
                if text.strip():
                    logger.info(f"pypdf: Extracted text from {extracted_pages}/{len(pdf_reader.pages)} pages")
                    return text.strip(), metadata
                        
        except Exception as e2:
            pypdf_error = str(e2)
            logger.error(f"pypdf also failed: {e2}")
        
        # If both methods failed or returned no text
        if not text.strip():
            error_msg = f"Both PDF readers failed to extract text. "
            if pypdf2_error:
                error_msg += f"PyPDF2: {pypdf2_error}. "
            if 'pypdf_error' in locals():
                error_msg += f"pypdf: {pypdf_error}. "
            error_msg += "This might be a scanned PDF, encrypted PDF, or the text might be embedded as images."
            raise Exception(error_msg)
        
        return text.strip(), metadata
    
    def _extract_docx_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX files.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            tuple[str, Dict[str, Any]]: Extracted text and metadata
        """
        try:
            if not os.path.exists(file_path):
                raise Exception(f"DOCX file not found: {file_path}")
            
            if os.path.getsize(file_path) == 0:
                raise Exception("DOCX file is empty (0 bytes)")
            
            doc = Document(file_path)
            text = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                text += f"\n--- Table {table_count} ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            metadata = {
                'paragraphs': paragraph_count,
                'tables': table_count,
                'method': 'python-docx'
            }
            
            if not text.strip():
                raise Exception("No text found in DOCX document")
            
            return text.strip(), metadata
            
        except Exception as e:
            if "No text found" in str(e):
                raise e
            raise Exception(f"Failed to extract text from DOCX: {e}")
    
    def _extract_txt_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from TXT files.
        
        Args:
            file_path (str): Path to TXT file
            
        Returns:
            tuple[str, Dict[str, Any]]: Extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise Exception(f"TXT file not found: {file_path}")
        
        if os.path.getsize(file_path) == 0:
            # Empty file is valid for TXT
            return "", {
                'encoding': 'utf-8',
                'lines': 0,
                'method': 'plain_text'
            }
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    
                metadata = {
                    'encoding': encoding,
                    'lines': len(text.splitlines()) if text else 0,
                    'method': 'plain_text'
                }
                
                return text.strip(), metadata
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                if encoding == encodings[-1]:  # Last encoding attempt
                    raise Exception(f"Failed to read text file: {e}")
                continue
        
        raise Exception("Failed to decode text file with any supported encoding")
    
    def get_supported_formats(self) -> list[str]:
        """
        Get list of supported file formats.
        
        Returns:
            list[str]: List of supported file extensions
        """
        return self.supported_formats.copy()
    
    def is_supported_format(self, filename: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            filename (str): Name of the file
            
        Returns:
            bool: True if format is supported
        """
        if not filename:
            return False
        return Path(filename).suffix.lower() in self.supported_formats